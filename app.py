# app.py (Final Version with Word & PDF Download)

import streamlit as st
import pandas as pd
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from collections import Counter
import requests
from bs4 import BeautifulSoup
import re
import time
import io

# New imports for Word and PDF generation
from docx import Document
from fpdf import FPDF

# --- 1. SET PAGE CONFIG ---
st.set_page_config(layout="wide", page_title="Skincare Data Analyst Tool")

# --- 2. DEFINE CONSTANTS AND HELPER FUNCTIONS ---

# Define a set of English stop words
STOP_WORDS = set([
    'a', 'about', 'above', 'after', 'again', 'against', 'all', 'am', 'an', 'and', 'any', 'are', "aren't", 'as', 'at',
    'be', 'because', 'been', 'before', 'being', 'below', 'between', 'both', 'but', 'by', 'can', "can't", 'cannot',
    'could', "couldn't", 'did', "didn't", 'do', 'does', "doesn't", 'doing', 'don', "don't", 'down', 'during', 'each',
    'few', 'for', 'from', 'further', 'had', "hadn't", 'has', "hasn't", 'have', "haven't", 'having', 'he', "he'd",
    "he'll", "he's", 'her', 'here', "here's", 'hers', 'herself', 'him', 'himself', 'his', 'how', "how's", 'i', "i'd",
    "i'll", "i'm", "i've", 'if', 'in', 'into', 'is', "isn't", 'it', "it's", 'its', 'itself', "let's", 'me', 'more',
    'most', "mustn't", 'my', 'myself', 'no', 'nor', 'not', 'of', 'off', 'on', 'once', 'only', 'or', 'other', 'ought',
    'our', 'ours', 'ourselves', 'out', 'over', 'own', 'same', 'shan', "shan't", 'she', "she'd", "she'll", "she's",
    'should', "shouldn't", 'so', 'some', 'such', 'than', 'that', "that's", 'the', 'their', 'theirs', 'them',
    'themselves', 'then', 'there', "there's", 'these', 'they', "they'd", "they'll", "they're", "they've", 'this',
    'those', 'through', 'to', 'too', 'under', 'until', 'up', 'very', 'was', "wasn't", 'we', "we'd", "we'll", "we're",
    "we've", 'were', "weren't", 'what', "what's", 'when', "when's", 'where', "where's", 'which', 'while', 'who',
    "who's", 'whom', 'why', "why's", 'with', 'won', "won't", 'would', "wouldn't", 'you', "you'd", "you'll", "you're",
    "you've", 'your', 'yours', 'yourself', 'yourselves'
])
STOP_WORDS.update(['skin', 'product', 'products', 'use', 'using', 'like', 'get', 'help', 'really', 'ive', 'im', 'would', 'routine', 'feel', 'look'])


# --- CORE ANALYSIS FUNCTIONS (UNCHANGED) ---
@st.cache_data
def analyze_product_mentions(text_data, product_list, claim_keywords, platform):
    sid = SentimentIntensityAnalyzer()
    product_mentions = []
    for text in text_data:
        for product in product_list:
            if re.search(r'\b' + re.escape(product.lower()) + r'\b', text.lower()):
                sentiment_score = sid.polarity_scores(text)['compound']
                if sentiment_score >= 0.05: sentiment = "Positive"
                elif sentiment_score <= -0.05: sentiment = "Negative"
                else: sentiment = "Mixed"
                found_claim = "General Skincare"
                for claim, keywords in claim_keywords.items():
                    if any(keyword in text.lower() for keyword in keywords):
                        found_claim = claim; break
                product_mentions.append({"Product Name": product, "Platform Mentioned On": platform, "Most Common Use or Claim": found_claim, "User Sentiment": sentiment})
    if not product_mentions: return pd.DataFrame()
    df = pd.DataFrame(product_mentions)
    summary = df.groupby(['Product Name', 'Platform Mentioned On', 'Most Common Use or Claim', 'User Sentiment']).size().reset_index(name='Approximate Number of Mentions')
    return summary.sort_values(by='Approximate Number of Mentions', ascending=False).reset_index(drop=True)

@st.cache_data
def get_trending_keywords(text_data):
    all_words = []
    for text in text_data:
        words = [word for word in re.findall(r'\b\w+\b', text.lower()) if word not in STOP_WORDS and len(word) > 2]
        all_words.extend(words)
    bigrams = list(zip(all_words, all_words[1:]))
    bigram_freq = Counter(bigrams)
    keyword_insights = []
    for phrase, _ in bigram_freq.most_common(5):
        keyword = ' '.join(phrase)
        explanation = "High frequency in discussions suggests strong user interest."
        if "acid" in keyword or "retin" in keyword: explanation = "Trending due to its popularity as a powerful active ingredient."
        if "sunscreen" in keyword or "spf" in keyword: explanation = "A cornerstone of skincare, consistently discussed."
        if "dark spots" in keyword: explanation = "A common concern users are actively seeking solutions for."
        keyword_insights.append({"Trending Keyword": keyword, "Reason for Trend": explanation})
    return pd.DataFrame(keyword_insights)

@st.cache_data
def trace_manufacturer(product_name):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
    query = f"{product_name} skincare official website"
    search_url = f"https://html.duckduckgo.com/html/?q={requests.utils.quote(query)}"
    try:
        response = requests.get(search_url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        link = soup.find('a', class_='result__a')
        if link and link['href']:
            clean_url = requests.utils.unquote(link['href']).split("uddg=")[-1]
            return clean_url, "Successfully found a likely official link."
        return "Not Found", "Could not automatically find a link."
    except requests.RequestException:
        return "Search Failed", "An error occurred during search."

# --- DOWNLOAD HELPER FUNCTIONS ---

def to_excel(df_dict):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for sheet_name, df in df_dict.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    return output.getvalue()

# NEW: Function to generate a Word document
def to_word(df_dict):
    document = Document()
    document.add_heading('Skincare Analysis Report', 0)

    for sheet_name, df in df_dict.items():
        if df.empty: continue
        document.add_heading(sheet_name, level=1)
        
        # Add a table to the document
        table = document.add_table(rows=1, cols=df.shape[1])
        table.style = 'Table Grid'
        
        # Add the header rows
        for i, column_name in enumerate(df.columns):
            table.cell(0, i).text = str(column_name)

        # Add the rest of the data frame
        for index, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        document.add_paragraph() # Add some space between tables

    # Save document to a byte stream
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# NEW: Function to generate a PDF document
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Skincare Analysis Report', 0, 1, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(5)

    def chapter_body(self, df):
        if df.empty:
            self.set_font('Arial', '', 10)
            self.cell(0, 10, "No data available for this section.", 0, 1)
            self.ln()
            return

        # Determine column widths
        col_width = self.w / (len(df.columns) + 0.5)  # Adjust divisor for padding
        
        # Header
        self.set_font('Arial', 'B', 8)
        for col in df.columns:
            self.cell(col_width, 10, col, 1, 0, 'C')
        self.ln()

        # Data
        self.set_font('Arial', '', 8)
        for index, row in df.iterrows():
            for item in row:
                self.multi_cell(col_width, 10, str(item), 1, 0, 'L')
            self.ln()
        self.ln(10) # Space after table


def to_pdf(df_dict):
    pdf = PDF(orientation='L', unit='mm', format='A4') # Landscape for more space
    pdf.add_page()
    
    for sheet_name, df in df_dict.items():
        pdf.chapter_title(sheet_name)
        pdf.chapter_body(df)
        
    # FPDF.output() returns bytes when the first argument is not a file name
    return pdf.output(dest='S').encode('latin-1')


# --- 3. BUILD THE APP LAYOUT ---
st.title("🧪 Skincare Data Analysis Assistant")
st.markdown("""
This tool helps analyze skincare discussions from online sources.
1.  **Configure** the product list and platform name in the sidebar.
2.  **Paste** your research data (e.g., from TikTok, Reddit) into the text box.
3.  **Click 'Analyze Data'** to see insights and download reports.
""")

# --- SIDEBAR FOR CONFIGURATION ---
with st.sidebar:
    st.header("⚙️ Configuration")
    DEFAULT_PRODUCTS = ["CeraVe", "The Ordinary", "La Roche-Posay", "Paula's Choice", "SkinCeuticals", "Drunk Elephant", "Supergoop", "EltaMD", "Tretinoin", "Kojic Acid", "Azelaic Acid", "Niacinamide", "Hyaluronic Acid", "Vitamin C Serum", "Anthelios"]
    CLAIM_KEYWORDS = {
        "Fades Dark Spots / Hyperpigmentation": ["hyperpigmentation", "dark spots", "pih", "melasma", "even tone", "brighter"],
        "Helps Acne / Breakouts": ["acne", "pimples", "breakouts", "clogged pores", "comedones"],
        "Anti-Aging / Wrinkles": ["anti-aging", "wrinkles", "fine lines", "aging", "retinol", "tretinoin"],
        "Hydration / Moisturizing": ["hydration", "dry skin", "moisture", "hydrating", "moisturizer", "plump"],
        "Sun Protection": ["sunscreen", "spf", "sun protection", "uv", "white cast"],
        "Improves Skin Texture": ["texture", "smooth", "bumpy", "keratosis pilaris", "kp", "pores"]
    }
    selected_products = st.multiselect("Select products to search for:", options=DEFAULT_PRODUCTS, default=DEFAULT_PRODUCTS[:7])
    platform_name = st.text_input("Platform Mentioned On:", value="TikTok / Online Forums")

# --- MAIN PAGE FOR INPUT AND OUTPUT ---
st.header("1. Paste Your Research Data")
default_text = "I've been using The Ordinary's Niacinamide for a month and my pores look smaller! Amazing for oily skin.\nHas anyone tried the La Roche-Posay Anthelios sunscreen? Worried about a white cast.\nUgh, the Drunk Elephant Vitamin C serum completely broke me out. A negative experience for me.\nThe CeraVe hydrating cleanser is my holy grail for dry skin. It's so gentle.\nLooking for a good hyperpigmentation fix. Someone suggested Kojic Acid soap.\nThis tretinoin journey is tough, my skin is peeling so much but my acne is getting better. It's a mixed bag."
user_text = st.text_area("Each new line should be a separate comment or post.", default_text, height=250)

if st.button("🚀 Analyze Data"):
    if not user_text.strip():
        st.warning("Please paste some data into the text box before analyzing.")
    elif not selected_products:
        st.warning("Please select at least one product to search for in the sidebar.")
    else:
        with st.spinner("Analyzing your data... This may take a moment."):
            text_lines = [line.strip() for line in user_text.strip().split('\n') if line.strip()]
            
            df_summary = analyze_product_mentions(text_lines, selected_products, CLAIM_KEYWORDS, platform_name)
            df_keywords = get_trending_keywords(text_lines)
            
            manufacturer_info = []
            if not df_summary.empty:
                top_2_products = df_summary['Product Name'].unique()[:2]
                for product in top_2_products:
                    link, note = trace_manufacturer(product)
                    manufacturer_info.append({"Product Name": product, "Official Link (Best Guess)": link, "Note": note})
                    time.sleep(0.5) 
            df_manufacturer = pd.DataFrame(manufacturer_info)

            st.success("Analysis Complete!")
            
            st.header("📊 Results")

            if df_summary.empty:
                st.info("No mentions of the selected products were found in the provided text.")
            else:
                st.subheader("1. Product Mention Summary")
                st.dataframe(df_summary, use_container_width=True)

                st.subheader("2. Trend & Keyword Insights")
                st.dataframe(df_keywords, use_container_width=True)
                
                st.subheader("3. Manufacturer Info (Basic Trace)")
                st.dataframe(df_manufacturer, use_container_width=True)

                # --- DOWNLOAD SECTION ---
                st.subheader("📥 Download Full Report")
                
                # Prepare data for download
                report_data_dict = {
                    'Product Mention Summary': df_summary,
                    'Trend & Keyword Insights': df_keywords,
                    'Manufacturer Info': df_manufacturer
                }
                
                # Generate files in memory
                excel_data = to_excel(report_data_dict)
                word_data = to_word(report_data_dict)
                pdf_data = to_pdf(report_data_dict)
                
                # Display download buttons in columns
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button(
                        label="📄 Download as Excel",
                        data=excel_data,
                        file_name="Skincare_Analysis_Report.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                with col2:
                    st.download_button(
                        label="📄 Download as Word",
                        data=word_data,
                        file_name="Skincare_Analysis_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with col3:
                    st.download_button(
                        label="📄 Download as PDF",
                        data=pdf_data,
                        file_name="Skincare_Analysis_Report.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
